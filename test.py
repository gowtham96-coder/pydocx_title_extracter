import re
import docx
from docx import Document
from docx.shared import Cm, Pt

headings_list = []
doc = docx.Document('ABHINAV_VANAMA-THESIS.docx')
for i in range(len(doc.paragraphs)):
    line = doc.paragraphs[i].text
    for run in doc.paragraphs[i].runs:
        if run.bold:
            pattern1 = '^Fig+(\.)'
            # pattern = '^Fig+(\.)+(\s)+(\d)+(\.)+(\d)'
            # pattern2 = '^Fig+(\.)+(\s)'
            # pattern3 = '^Fig+(\.)+(\d)+(\.)+(\d)'
            result = re.match(pattern1, line)
            # result1 = re.search(pattern2, line)
            if result:
                headings_list.append(line)
            # if result1:
            #     headings_list.append(line)
Fig = set(headings_list)
# print(Fig)

Dict = dict((x.strip(), y.strip())
            for x, y in (element.split(' ', 1)
                         for element in Fig))
# printing result
sorted_dict = {}
sorted_keys = sorted(Dict.keys())  # [1, 3, 2]
for w in sorted_keys:
    sorted_dict[w] = Dict[w]

print(sorted_dict)
doc = docx.Document()

# Add a Title to the document
doc.add_heading('LIST OF FIGURES', 0)

# Table data in a form of list
# Creating a table object
table = doc.add_table(rows=1, cols=3)

# Adding heading in the 1st row of the table
row = table.rows[0].cells
row[0].text = 'Figure No'
row[1].text = 'Title'
row[2].text = 'Page No'

# Adding data from the list to the table
for chapter, title in sorted_dict.items():

    # Adding a row and then adding data in it.
    row = table.add_row().cells
    # Converting id to string as table can only take string input
    row[0].text = str(chapter)
    row[1].text = title
    # row[2].text = " "
table.style = 'Colorful List'
# Now save the document to a location
doc.save('Content_Test.docx')
