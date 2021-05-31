import re
import docx


class Context:
    def __init__(gow):
        # import os
        gow.headings_list = []
        gow.fig_list = []
        gow.table_list = []
        gow.file = docx.Document()
        # File = os.environ('file_path')

        gow.doc = docx.Document('ABHINAV_VANAMA-THESIS.docx')

    def headings(gow):
        for i in range(len(gow.doc.paragraphs)):
            line = gow.doc.paragraphs[i].text
            for run in gow.doc.paragraphs[i].runs:
                if run.bold:
                    pattern = '^(\d)+(\.)+(\d)+(\s)'
                    pattern2 = '^(\d)+(\.)+(\d)+(\.)+(\d)+(\s)'
                    result = re.match(pattern, line)
                    subheading = re.match(pattern2, line)
                    if result:
                        gow.headings_list.append(line)
                    if subheading:
                        gow.headings_list.append(line)

        gow.heading_dict = dict((x.strip(), y.strip())
                                for x, y in (element.split(' ', 1)
                                             for element in gow.headings_list))

    def figures(gow):
        for i in range(len(gow.doc.paragraphs)):
            line = gow.doc.paragraphs[i].text
            for run in gow.doc.paragraphs[i].runs:
                if run.bold:
                    pattern1 = '^Fig+(\.)'
                    result = re.match(pattern1, line)
                    if result:
                        line = line.replace('Fig.', "")
                        gow.fig_list.append(line)
        Fig = set(gow.fig_list)
        Dict = dict((x.strip(), y.strip())
                    for x, y in (element.split(' ', 1)
                                 for element in Fig))
        gow.sorted_fig = {}
        sorted_keys = sorted(Dict.keys())  # [1, 3, 2]
        for w in sorted_keys:
            gow.sorted_fig[w] = Dict[w]

    def tables(gow):
        for i in range(len(gow.doc.paragraphs)):
            line = gow.doc.paragraphs[i].text
            for run in gow.doc.paragraphs[i].runs:
                if run.bold:
                    tab_pattern = '^Table'
                    tab_result = re.match(tab_pattern, line)
                    if tab_result:
                        line = line.replace('Table ', "")
                        gow.table_list.append(line)
        table = set(gow.table_list)
        tab_dict = dict((x.strip(), y.strip())
                    for x, y in (element.split(' ', 1)
                                 for element in table))
        gow.sorted_tables = {}
        sorted_keys = sorted(tab_dict.keys())  # [1, 3, 2]
        for w in sorted_keys:
            gow.sorted_tables[w] = tab_dict[w]

    def prepare_docs(gow, Dict, title, name):
        # Add a Title to the document
        gow.file.add_heading(title, 0)
        # Table data in a form of list
        # Creating a table object
        table = gow.file.add_table(rows=1, cols=3)
        # Adding heading in the 1st row of the table
        row = table.rows[0].cells
        row[0].text = name
        row[1].text = 'Title'
        row[2].text = 'Page No'
        # Adding data from the list to the table
        for chapter, title in Dict.items():
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(chapter)
            row[1].text = title
            # row[2].text = " "
        table.style = 'Colorful List'
        # Now save the document to a location
        gow.file.save('Content_Test.docx')


c1 = Context()
c1.headings()
print(c1.heading_dict)
c1.prepare_docs(c1.heading_dict,"CONTENT","Chapter")
c1.figures()
print(c1.sorted_fig)
c1.prepare_docs(c1.sorted_fig,"\nLIST OF FIGURES","Figure No")
c1.tables()
print(c1.sorted_tables)
c1.prepare_docs(c1.sorted_tables,"\nLIST OF TABLES","Table No")
