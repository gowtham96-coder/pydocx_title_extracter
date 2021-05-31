import re
import docx
from docx import Document
from docx.shared import Cm, Pt
headings_list = []
doc = docx.Document('ABHINAV_VANAMA-THESIS.docx')
for i in range(len(doc.paragraphs)):
    line = doc.paragraphs[i].text
    for run in doc.paragraphs[i].runs:
        if run.bold:
            pattern = '^(\d)+(\.)+(\d)+(\s)'
            pattern2 = '^(\d)+(\.)+(\d)+(\.)+(\d)+(\s)'
            result = re.match(pattern, line)
            subheading = re.match(pattern2, line)
            if result:
                headings_list.append(line)
            if subheading:
                headings_list.append(line)

Dict = dict((x.strip(), y.strip())
            for x, y in (element.split(' ', 1)
                         for element in headings_list))

print(Dict)


doc = docx.Document()

# Add a Title to the document
doc.add_heading('CONTENT', 0)

# Table data in a form of list
# Creating a table object
table = doc.add_table(rows=1, cols=3)

# Adding heading in the 1st row of the table
row = table.rows[0].cells
row[0].text = 'Chapter'
row[1].text = 'Title'
row[2].text = 'Page No'

# Adding data from the list to the table
for chapter, title in Dict.items():

    # Adding a row and then adding data in it.
    row = table.add_row().cells
    # Converting id to string as table can only take string input
    row[0].text = str(chapter)
    row[1].text = title
    # row[2].text = " "
table.style = 'Colorful List'
# Now save the document to a location
doc.save('Content_Test.docx')