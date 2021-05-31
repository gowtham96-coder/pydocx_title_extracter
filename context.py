import re
import docx
import os
from werkzeug.utils import secure_filename
from flask import Flask, flash, request, redirect, send_file, render_template


class Context:
    def __init__(gow):
        # import os
        gow.headings_list = []
        gow.fig_list = []
        gow.table_list = []
        gow.file = docx.Document()

    def headings(gow, path):
        gow.doc = docx.Document(path)
        for i in range(len(gow.doc.paragraphs)):
            line = gow.doc.paragraphs[i].text
            for run in gow.doc.paragraphs[i].runs:
                if run.bold:
                    pattern = '^(\d)+(\.)+(\d)+(\s)'
                    pattern2 = '^(\d)+(\.)+(\d)+(\.)+(\d)+(\s)'
                    result = re.match(pattern, line)
                    subheading = re.match(pattern2, line)
                    if result:
                        gow.headings_list.append(line)
                    if subheading:
                        gow.headings_list.append(line)

        gow.heading_dict = dict((x.strip(), y.strip())
                                for x, y in (element.split(' ', 1)
                                             for element in gow.headings_list))

    def figures(gow, path):
        gow.doc = docx.Document(path)
        for i in range(len(gow.doc.paragraphs)):
            line = gow.doc.paragraphs[i].text
            for run in gow.doc.paragraphs[i].runs:
                if run.bold:
                    pattern1 = '^Fig+(\.)'
                    result = re.match(pattern1, line)
                    if result:
                        line = line.replace('Fig.', "")
                        gow.fig_list.append(line)
        Fig = set(gow.fig_list)
        Dict = dict((x.strip(), y.strip())
                    for x, y in (element.split(' ', 1)
                                 for element in Fig))
        gow.sorted_fig = {}
        sorted_keys = sorted(Dict.keys())  # [1, 3, 2]
        for w in sorted_keys:
            gow.sorted_fig[w] = Dict[w]

    def tables(gow, path):
        gow.doc = docx.Document(path)
        for i in range(len(gow.doc.paragraphs)):
            line = gow.doc.paragraphs[i].text
            for run in gow.doc.paragraphs[i].runs:
                if run.bold:
                    tab_pattern = '^Table'
                    tab_result = re.match(tab_pattern, line)
                    if tab_result:
                        line = line.replace('Table ', "")
                        gow.table_list.append(line)
        table = set(gow.table_list)
        tab_dict = dict((x.strip(), y.strip())
                        for x, y in (element.split(' ', 1)
                                     for element in table))
        gow.sorted_tables = {}
        sorted_keys = sorted(tab_dict.keys())  # [1, 3, 2]
        for w in sorted_keys:
            gow.sorted_tables[w] = tab_dict[w]

    def prepare_docs(gow, Dict, title, name):
        # Add a Title to the document
        gow.file.add_heading(title, 0)
        # Table data in a form of list
        # Creating a table object
        table = gow.file.add_table(rows=1, cols=3)
        # Adding heading in the 1st row of the table
        row = table.rows[0].cells
        row[0].text = name
        row[1].text = 'Title'
        row[2].text = 'Page No'
        # Adding data from the list to the table
        for chapter, title in Dict.items():
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(chapter)
            row[1].text = title
            # row[2].text = " "
        table.style = 'Colorful List'
        # Now save the document to a location
        gow.file.save('Content_word.docx')

    def url_call(gow):
        UPLOAD_FOLDER = 'C:/Users/gowth/PycharmProjects/contextcode/uploads/'

        # app = Flask(__name__)
        gow.app = Flask(__name__, template_folder='templates')
        gow.app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

        # app.secret_key = '123'

        # Upload API
        @gow.app.route('/', methods=['GET', 'POST'])
        def upload_file():
            if request.method == 'POST':
                # check if the post request has the file part
                if 'file' not in request.files:
                    print('no file')
                    return redirect(request.url)
                file = request.files['file']
                # if user does not select file, browser also
                # submit a empty part without filename
                if file.filename == '':
                    print('no filename')
                    return redirect(request.url)
                else:
                    filename = secure_filename(file.filename)
                    file.save(os.path.join(gow.app.config['UPLOAD_FOLDER'], filename))
                    print("saved file successfully")
                    # send file name as parameter to downlad
                    process_file(os.path.join(gow.app.config['UPLOAD_FOLDER'], filename))
                    return redirect('/downloadfile/' + filename)

            return render_template('upload_file.html')

        def process_file(path):
            gow.headings(path)
            gow.figures(path)
            gow.tables(path)
            gow.prepare_docs(gow.heading_dict, "CONTENT", "Chapter")
            gow.prepare_docs(gow.sorted_fig, "\nLIST OF FIGURES", "Figure No")
            gow.prepare_docs(gow.sorted_tables, "\nLIST OF TABLES", "Table No")

        # Download API
        @gow.app.route("/downloadfile/<filename>", methods=['GET'])
        def download_file(filename):
            return render_template('download.html', value=filename)

        @gow.app.route('/return-files/<filename>')
        def return_files_tut(filename):
            # file_path = UPLOAD_FOLDER + filename
            file_path = 'C:/Users/gowth/PycharmProjects/contextcode/Content_word.docx'
            return send_file(file_path, as_attachment=True, attachment_filename='')


if __name__ == "__main__":
    c1 = Context()
    c1.url_call()
    c1.app.run(debug=True)
